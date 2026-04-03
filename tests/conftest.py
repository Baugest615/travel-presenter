"""共用 fixtures — Travel Presenter 測試基礎"""
import json
import tempfile
from pathlib import Path

import pytest
from docx import Document

from travel_presenter.models import (
    TripData, DayItinerary, Flight, Hotel, Meals, MeetingPoint,
)


# ── 路徑 ─────────────────────────────────────────────────
FIXTURES_DIR = Path(__file__).parent / "fixtures"
EXAMPLES_DIR = Path(__file__).parent.parent / "examples"


# ── 最小合法 TripData ─────────────────────────────────────
@pytest.fixture
def minimal_trip() -> TripData:
    """只有必填欄位的最小合法 TripData"""
    return TripData(
        title="測試行程",
        date_range="2026/04/01 — 04/03",
        destination="東京",
    )


@pytest.fixture
def full_trip() -> TripData:
    """包含所有欄位的完整 TripData"""
    return TripData(
        title="2026 北海道七日遊",
        subtitle="破冰船・雪上樂園",
        company="測試旅行社",
        date_range="2026/03/15 — 03/21",
        destination="北海道",
        quote="在最寒冷的季節，遇見最溫暖的風景",
        quote_en="In the coldest season, we find the warmest scenery.",
        theme="soft-cream",
        flights=[
            Flight(
                direction="departure",
                airline="長榮航空",
                flight_number="BR116",
                date="2026/03/15",
                departure_airport="TPE",
                departure_time="09:30",
                arrival_airport="CTS",
                arrival_time="14:05",
            ),
            Flight(
                direction="return",
                airline="長榮航空",
                flight_number="BR115",
                date="2026/03/21",
                departure_airport="CTS",
                departure_time="15:20",
                arrival_airport="TPE",
                arrival_time="19:05",
            ),
        ],
        meeting_point=MeetingPoint(
            time="2026/03/15 07:00",
            location="桃園國際機場 第二航廈",
        ),
        days=[
            DayItinerary(
                day=1,
                date="3/15（日）",
                title="森林精靈露台",
                title_en="Ningle Terrace, Furano",
                route="桃園機場 → 新千歲機場 → 富良野",
                meals=Meals(breakfast="機上餐食", lunch="拉麵定食", dinner="飯店自助餐"),
                hotel=Hotel(name="富良野王子大飯店", area="FURANO"),
            ),
            DayItinerary(
                day=2,
                date="3/16（一）",
                title="旭山動物園",
                route="富良野 → 旭川",
                meals=Meals(breakfast="飯店早餐", lunch="味噌拉麵", dinner="成吉思汗烤肉"),
                hotel=Hotel(name="旭川 JR Inn"),
            ),
        ],
        hotels=[
            Hotel(name="富良野王子大飯店", area="FURANO", nights=[1]),
            Hotel(name="旭川 JR Inn", nights=[2]),
        ],
        notes=["請攜帶保暖衣物", "護照效期需六個月以上"],
    )


@pytest.fixture
def sample_json_path() -> Path:
    """範例 JSON 檔案路徑"""
    return EXAMPLES_DIR / "hokkaido_sample.json"


# ── DOCX fixture 工廠 ─────────────────────────────────────
@pytest.fixture
def make_docx(tmp_path):
    """產生自訂內容的 DOCX 檔案"""
    def _make(tables=None, paragraphs=None) -> Path:
        doc = Document()

        # 先加段落
        if paragraphs:
            for text in paragraphs:
                doc.add_paragraph(text)

        # 再加表格
        if tables:
            for tbl_spec in tables:
                rows = tbl_spec.get("rows", [])
                if not rows:
                    continue
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                for i, row_data in enumerate(rows):
                    for j, cell_text in enumerate(row_data):
                        if j < ncols:
                            table.rows[i].cells[j].text = cell_text

        path = tmp_path / "test.docx"
        doc.save(str(path))
        return path

    return _make


@pytest.fixture
def make_json(tmp_path):
    """產生自訂內容的 JSON 檔案"""
    def _make(data: dict, filename: str = "test.json") -> Path:
        path = tmp_path / filename
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return path

    return _make
